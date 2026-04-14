from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# ---- Styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.name = 'Times New Roman'
    heading_style.font.color.rgb = RGBColor(0, 0, 0)

# ---- Title Page ----
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Software Requirement Specification (SRS)')
run.bold = True
run.font.size = Pt(26)
run.font.name = 'Times New Roman'

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('REPO 2D \u2013 Dark Retrieval Game')
run.bold = True
run.font.size = Pt(20)
run.font.name = 'Times New Roman'

doc.add_paragraph()

tech = doc.add_paragraph()
tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tech.add_run('Developed using Godot Engine 4.6.1 (Mono)')
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

for _ in range(3):
    doc.add_paragraph()

info_data = [
    ('Project Title', 'REPO 2D \u2013 Dark Retrieval Game'),
    ('Technology Used', 'Godot Engine 4.6.1, GDScript, GLSL Shaders'),
    ('Platform', 'Windows Desktop (PC)'),
    ('Department', 'Computer Science & Engineering'),
    ('Semester', '6th Semester'),
    ('Date', '25/03/2026'),
]
table = doc.add_table(rows=len(info_data), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (key, val) in enumerate(info_data):
    row = table.rows[i]
    row.cells[0].text = key
    row.cells[1].text = val

doc.add_page_break()

# ---- TOC ----
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction', '2. Problem Statement', '3. Objectives', '4. Scope',
    '5. Functional Requirements', '6. Non-Functional Requirements',
    '7. System Architecture', '8. Technology Stack', '9. Module Description',
    '10. Database Design', '11. User Interface Design',
    '12. Testing Strategy', '13. Future Enhancements',
    '14. Conclusion', '15. References',
]
for item in toc_items:
    doc.add_paragraph(item)

doc.add_page_break()

# ---- 1. Introduction ----
doc.add_heading('1. Introduction', level=1)
doc.add_paragraph(
    'REPO 2D \u2013 Dark Retrieval Game is a top-down 2D horror-survival game developed using the '
    'Godot Engine 4.6.1 (Mono Edition). The game is inspired by the popular horror game "REPO" '
    'and reimagines its core mechanics in a 2D environment with advanced procedural graphics, '
    'dynamic lighting, and AI-driven enemy behaviour.')
doc.add_paragraph(
    'The player assumes the role of a "Repo Agent" who must enter dark, abandoned facilities to '
    'retrieve valuable items and meet a monetary quota. The player is equipped with a flashlight '
    '(limited battery) and must navigate procedurally generated maze-like levels while avoiding '
    'shadow creatures that patrol, detect, and chase the player.')
doc.add_paragraph(
    'This document outlines the complete Software Requirement Specification (SRS) for the project, '
    'covering functional and non-functional requirements, system architecture, module descriptions, '
    'and the overall technical design of the game.')

# ---- 2. Problem Statement ----
doc.add_heading('2. Problem Statement', level=1)
doc.add_paragraph(
    'Most horror games available today are 3D-based and require high-end hardware (GPU, RAM) to run '
    'effectively. This creates a barrier for players with low-specification PCs. Additionally, '
    'many 2D horror games lack advanced graphical effects such as dynamic lighting, real-time shadows, '
    'post-processing shaders, and intelligent enemy AI, resulting in a less immersive experience.')
doc.add_paragraph(
    'There is a need for a lightweight yet visually advanced 2D horror game that can run on any '
    'Windows PC without requiring dedicated GPUs, while still delivering an engaging and atmospheric '
    'horror-survival experience with progressive difficulty.')

# ---- 3. Objectives ----
doc.add_heading('3. Objectives', level=1)
objectives = [
    'To develop a 2D top-down horror-survival game using Godot Engine 4.6.1.',
    'To implement procedural level generation ensuring unique gameplay each session.',
    'To create advanced 2D lighting and shadow systems using PointLight2D and LightOccluder2D.',
    'To design intelligent enemy AI with state-machine-based behaviour (Patrol, Investigate, Chase).',
    'To implement post-processing visual effects using custom GLSL shaders (Vignette, Fear Overlay).',
    'To build a progressive difficulty system spanning 5 levels with increasing challenges.',
    'To ensure the game runs responsively on any screen resolution and aspect ratio.',
    'To export the game as a standalone Windows executable (.exe) for easy distribution.',
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

# ---- 4. Scope ----
doc.add_heading('4. Scope', level=1)
doc.add_heading('4.1 In Scope', level=2)
for item in [
    'Single-player horror-survival gameplay with 5 progressive levels.',
    'Procedurally generated maze layouts using grid-based room algorithms.',
    'Dynamic 2D lighting system with flashlight mechanics and battery management.',
    'Enemy AI with line-of-sight detection using RayCast2D.',
    'Item collection system with quota tracking.',
    'HUD displaying level number, quota progress, battery status, and survival time.',
    'Post-processing shaders for vignette and fear distortion effects.',
    'Standalone Windows desktop export (.exe).',
    'Source code hosting on GitHub for version control.',
]:
    doc.add_paragraph(item, style='List Bullet')
doc.add_heading('4.2 Out of Scope', level=2)
for item in [
    'Multiplayer or online co-op functionality.',
    'Mobile platform (Android/iOS) support.',
    'Audio/Sound effects implementation (planned for future versions).',
    'Save/Load game progress system.',
]:
    doc.add_paragraph(item, style='List Bullet')

# ---- 5. Functional Requirements ----
doc.add_heading('5. Functional Requirements', level=1)
fr_table = doc.add_table(rows=1, cols=4)
fr_table.style = 'Table Grid'
hdr = fr_table.rows[0].cells
for i, h in enumerate(['FR ID', 'Requirement', 'Description', 'Priority']):
    hdr[i].text = h
    for r in hdr[i].paragraphs[0].runs: r.bold = True
for fid, req, desc, pri in [
    ('FR-01', 'Player Movement', 'Player can move using WASD keys in a top-down 2D environment.', 'High'),
    ('FR-02', 'Flashlight Control', 'Player aims flashlight with mouse; toggle on/off with F key.', 'High'),
    ('FR-03', 'Battery System', 'Flashlight consumes battery over time; disables when depleted.', 'High'),
    ('FR-04', 'Sprint Mechanic', 'Player can sprint using Shift key; consumes stamina.', 'Medium'),
    ('FR-05', 'Item Collection', 'Player collects glowing crystal items to fill quota.', 'High'),
    ('FR-06', 'Quota System', 'Each level has a target quota; level completes when met.', 'High'),
    ('FR-07', 'Enemy Patrol AI', 'Enemies patrol between random waypoints in the maze.', 'High'),
    ('FR-08', 'Enemy Detection', 'Enemies detect player using RayCast2D line-of-sight.', 'High'),
    ('FR-09', 'Enemy Chase AI', 'Detected enemies chase player at increased speed.', 'High'),
    ('FR-10', 'Fear Mechanic', 'Proximity to enemies increases fear level; 100% = game over.', 'Medium'),
    ('FR-11', 'Procedural Levels', 'Maze layout is randomly generated each game session.', 'High'),
    ('FR-12', 'Progressive Difficulty', '5 levels with increasing maze size, quota, and enemies.', 'High'),
    ('FR-13', 'Main Menu', 'Start screen with Start Shift and Quit buttons.', 'Medium'),
    ('FR-14', 'Game Over Screen', 'Displays win/lose status with Next Level or Return to Menu.', 'Medium'),
    ('FR-15', 'HUD Display', 'Shows Level, Quota, Battery, and Time on screen.', 'High'),
]:
    row = fr_table.add_row().cells
    row[0].text = fid; row[1].text = req; row[2].text = desc; row[3].text = pri

# ---- 6. Non-Functional Requirements ----
doc.add_heading('6. Non-Functional Requirements', level=1)
nfr_table = doc.add_table(rows=1, cols=3)
nfr_table.style = 'Table Grid'
hdr = nfr_table.rows[0].cells
for i, h in enumerate(['NFR ID', 'Requirement', 'Description']):
    hdr[i].text = h
    for r in hdr[i].paragraphs[0].runs: r.bold = True
for nid, req, desc in [
    ('NFR-01', 'Performance', 'Game must maintain 60 FPS on systems with integrated graphics.'),
    ('NFR-02', 'Responsiveness', 'UI and gameplay must adapt to any screen resolution.'),
    ('NFR-03', 'Portability', 'Exportable as standalone .exe without requiring Godot.'),
    ('NFR-04', 'Usability', 'Controls must be intuitive (WASD + Mouse).'),
    ('NFR-05', 'Maintainability', 'Codebase must be modular with separate scripts per entity.'),
    ('NFR-06', 'Scalability', 'Level generation supports arbitrary grid sizes.'),
    ('NFR-07', 'Visual Quality', 'Advanced 2D lighting, shaders, and procedural animations.'),
    ('NFR-08', 'Reliability', 'Scene transitions use call_deferred to prevent crashes.'),
]:
    row = nfr_table.add_row().cells
    row[0].text = nid; row[1].text = req; row[2].text = desc

# ---- 7. System Architecture ----
doc.add_heading('7. System Architecture', level=1)
doc.add_paragraph(
    'The game follows a modular scene-based architecture inherent to the Godot Engine. '
    'Each game entity (Player, Enemy, Item, HUD, Level) is a separate scene with its own script.')
doc.add_heading('7.1 Architecture Overview', level=2)
arch_table = doc.add_table(rows=1, cols=3)
arch_table.style = 'Table Grid'
hdr = arch_table.rows[0].cells
for i, h in enumerate(['Layer', 'Components', 'Responsibility']):
    hdr[i].text = h
    for r in hdr[i].paragraphs[0].runs: r.bold = True
for layer, comp, resp in [
    ('Global', 'GameManager (Autoload)', 'State management, signals, scene transitions'),
    ('Scene', 'MainMenu, GameLevel, GameOver', 'Scene lifecycle and UI flow'),
    ('Entity', 'Player, Enemy, CollectibleItem', 'Gameplay logic, physics, AI'),
    ('UI', 'HUD (CanvasLayer)', 'Quota, Battery, Timer, Shader overlays'),
    ('Visual', 'Shaders (Vignette, Fear)', 'Post-processing effects'),
]:
    row = arch_table.add_row().cells
    row[0].text = layer; row[1].text = comp; row[2].text = resp
doc.add_heading('7.2 Scene Flow', level=2)
doc.add_paragraph('Main Menu -> Game Level (Level 1-5) -> Game Over Screen -> Next Level / Main Menu')

# ---- 8. Technology Stack ----
doc.add_heading('8. Technology Stack', level=1)
tech_table = doc.add_table(rows=1, cols=2)
tech_table.style = 'Table Grid'
hdr = tech_table.rows[0].cells
for i, h in enumerate(['Component', 'Technology']):
    hdr[i].text = h
    for r in hdr[i].paragraphs[0].runs: r.bold = True
for comp, t in [
    ('Game Engine', 'Godot Engine 4.6.1 Stable (Mono Edition)'),
    ('Programming Language', 'GDScript'),
    ('Shader Language', 'GLSL (Godot Shading Language)'),
    ('Version Control', 'Git + GitHub'),
    ('Target Platform', 'Windows Desktop (64-bit)'),
    ('Build System', 'Godot Export Templates'),
]:
    row = tech_table.add_row().cells
    row[0].text = comp; row[1].text = t

# ---- 9. Module Description ----
doc.add_heading('9. Module Description', level=1)
modules = [
    ('9.1 GameManager Module', 'game_manager.gd',
     'Global singleton managing game state, level tracking (1-5), quota, battery, fear, and scene transitions.'),
    ('9.2 Player Module', 'player.gd',
     'CharacterBody2D with WASD movement, sprint, flashlight aiming (mouse), battery drain, and _draw() rendering.'),
    ('9.3 Enemy Module', 'enemy.gd',
     'CharacterBody2D with FSM (PATROL, INVESTIGATE, CHASE, ATTACK). Uses RayCast2D for line-of-sight detection.'),
    ('9.4 Collectible Item Module', 'collectible_item.gd',
     'Area2D pickup with bobbing animation, glowing light, and light burst effect on collection.'),
    ('9.5 Level Generator Module', 'game_level.gd',
     'Procedural maze generator. Grid scales with level (Level 1=4x4, Level 5=8x8). Creates walls with LightOccluder2D.'),
    ('9.6 HUD Module', 'hud.gd',
     'CanvasLayer showing Level, Quota, Battery bar, Timer. Applies Vignette and Fear shaders.'),
    ('9.7 Main Menu Module', 'main_menu.gd',
     'Title screen with START SHIFT and QUIT buttons.'),
    ('9.8 Game Over Module', 'game_over.gd',
     'Results screen with contextual messages and NEXT LEVEL / RETURN TO MENU / VICTORY buttons.'),
    ('9.9 Shader Modules', 'vignette.gdshader, fear_overlay.gdshader',
     'GLSL shaders for vignette darkening and fear-based screen distortion with red tint and scanline glitch.'),
]
for heading, filename, description in modules:
    doc.add_heading(heading, level=2)
    p = doc.add_paragraph()
    p.add_run('File: ').bold = True
    p.add_run(filename)
    doc.add_paragraph(description)

# ---- 10. Database Design ----
doc.add_heading('10. Database Design', level=1)
doc.add_paragraph(
    'This project does not use any external database. All game state is managed in-memory '
    'by the GameManager autoload singleton. Future versions may implement local JSON-based saves.')

# ---- 11. UI Design ----
doc.add_heading('11. User Interface Design', level=1)
doc.add_heading('11.1 Main Menu Screen', level=2)
doc.add_paragraph('Dark background, large red title "REPO 2D" (96pt), subtitle, START SHIFT / QUIT buttons.')
doc.add_heading('11.2 In-Game HUD', level=2)
doc.add_paragraph('Top-left: Level indicator (golden, 48pt), Quota counter, Battery bar, Timer. Full-screen: Vignette + Fear shaders.')
doc.add_heading('11.3 Game Over Screen', level=2)
doc.add_paragraph('Contextual title (LEVEL X CLEARED / DIED IN THE DARK / YOU ESCAPED), stats, action button.')

# ---- 12. Testing Strategy ----
doc.add_heading('12. Testing Strategy', level=1)
test_table = doc.add_table(rows=1, cols=4)
test_table.style = 'Table Grid'
hdr = test_table.rows[0].cells
for i, h in enumerate(['Test ID', 'Test Case', 'Expected Result', 'Status']):
    hdr[i].text = h
    for r in hdr[i].paragraphs[0].runs: r.bold = True
for tid, tc, exp, st in [
    ('TC-01', 'Player WASD movement', 'Player moves in 4 directions', 'Passed'),
    ('TC-02', 'Flashlight toggle (F)', 'Light on/off; battery stops draining', 'Passed'),
    ('TC-03', 'Sprint (Shift)', 'Speed increases; stamina depletes', 'Passed'),
    ('TC-04', 'Item collection', 'Quota increases; light burst plays', 'Passed'),
    ('TC-05', 'Enemy patrol', 'Enemies move between waypoints', 'Passed'),
    ('TC-06', 'Enemy chase', 'Enemy pursues player on detection', 'Passed'),
    ('TC-07', 'Fear system', 'Fear increases near enemies; distortion activates', 'Passed'),
    ('TC-08', 'Level completion', 'Quota met triggers Game Over with NEXT LEVEL', 'Passed'),
    ('TC-09', 'Level 5 victory', 'YOU ESCAPED message after Level 5', 'Passed'),
    ('TC-10', 'Screen responsiveness', 'Game scales on different monitors', 'Passed'),
    ('TC-11', 'Scene transitions', 'No crash (call_deferred fix)', 'Passed'),
    ('TC-12', 'Procedural levels', 'Unique maze each playthrough', 'Passed'),
]:
    row = test_table.add_row().cells
    row[0].text = tid; row[1].text = tc; row[2].text = exp; row[3].text = st

# ---- 13. Future Enhancements ----
doc.add_heading('13. Future Enhancements', level=1)
for item in [
    'Add immersive sound effects and background music.',
    'Implement save/load system using JSON.',
    'Add multiplayer co-op mode (2-4 players).',
    'Introduce new enemy types with varied AI.',
    'Implement shop/upgrade system between levels.',
    'Port to Android and Web (HTML5) platforms.',
    'Add procedurally generated story elements.',
    'Implement leaderboards and achievements.',
]:
    doc.add_paragraph(item, style='List Bullet')

# ---- 14. Conclusion ----
doc.add_heading('14. Conclusion', level=1)
doc.add_paragraph(
    'REPO 2D successfully demonstrates the development of an advanced 2D horror-survival game '
    'using Godot Engine 4.6.1. The project showcases procedural level generation, dynamic 2D lighting '
    'with real-time shadows, intelligent enemy AI using finite state machines, and custom GLSL '
    'shader-based post-processing effects. The game features a progressive 5-level campaign where '
    'difficulty dynamically scales. All graphics are rendered programmatically using _draw() API, '
    'eliminating dependency on external sprites while maintaining a visually impressive experience.')

# ---- 15. References ----
doc.add_heading('15. References', level=1)
for i, ref in enumerate([
    'Godot Engine Official Documentation - https://docs.godotengine.org/en/stable/',
    'GDScript Language Reference - https://docs.godotengine.org/en/stable/tutorials/scripting/gdscript/',
    'Godot Shading Language - https://docs.godotengine.org/en/stable/tutorials/shaders/',
    'Godot 2D Lighting - https://docs.godotengine.org/en/stable/tutorials/2d/2d_lights_and_shadows.html',
    'REPO (Original Game) - Inspiration for game concept and mechanics.',
    'GitHub Repository - https://github.com/thanetraj/RepoGame2D',
], 1):
    doc.add_paragraph(f'[{i}] {ref}')

# ---- Save ----
output_path = r'D:\Thanet Projects\RepoGame2D\SRS_REPO_2D_Game.docx'
doc.save(output_path)
print(f'SRS Document saved to: {output_path}')
